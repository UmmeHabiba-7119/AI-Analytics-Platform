from io import BytesIO

import pandas as pd

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import (
    getSampleStyleSheet,
    ParagraphStyle
)
from reportlab.lib.units import inch
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
    Image
)

from docx import Document
from docx.shared import Inches


def create_pdf_report(
    metadata,
    filters,
    question,
    result_df,
    narrative,
    chart_caption="",
    chart_png=None
):
    """
    Create a formatted PDF report in memory.
    """

    buffer = BytesIO()

    document = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=40,
        leftMargin=40,
        topMargin=40,
        bottomMargin=40
    )

    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        "ReportTitle",
        parent=styles["Title"],
        alignment=TA_CENTER,
        spaceAfter=18
    )

    story = []

    story.append(
        Paragraph(
            "AI Analytics Report",
            title_style
        )
    )

    story.append(
        Paragraph(
            "<b>Dataset Metadata</b>",
            styles["Heading2"]
        )
    )

    metadata_rows = [
        ["Item", "Value"]
    ]

    for key, value in metadata.items():
        metadata_rows.append(
            [str(key), str(value)]
        )

    metadata_table = Table(
        metadata_rows,
        colWidths=[2.0 * inch, 4.5 * inch]
    )

    metadata_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold")
            ]
        )
    )

    story.append(metadata_table)
    story.append(Spacer(1, 14))

    story.append(
        Paragraph(
            "<b>Applied Filters</b>",
            styles["Heading2"]
        )
    )

    filter_rows = [
        ["Filter", "Selected Value"]
    ]

    for key, value in filters.items():
        filter_rows.append(
            [str(key), str(value)]
        )

    filter_table = Table(
        filter_rows,
        colWidths=[2.0 * inch, 4.5 * inch]
    )

    filter_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold")
            ]
        )
    )

    story.append(filter_table)
    story.append(Spacer(1, 14))

    story.append(
        Paragraph(
            "<b>User Question</b>",
            styles["Heading2"]
        )
    )

    story.append(
        Paragraph(
            str(question),
            styles["BodyText"]
        )
    )

    story.append(Spacer(1, 12))

    story.append(
        Paragraph(
            "<b>Query Result</b>",
            styles["Heading2"]
        )
    )

    display_df = result_df.head(20).copy()

    table_data = [
        [str(column) for column in display_df.columns]
    ]

    for _, row in display_df.iterrows():
        table_data.append(
            [
                str(value)[:40]
                for value in row.tolist()
            ]
        )

    if table_data:
        result_table = Table(
            table_data,
            repeatRows=1
        )

        result_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                    ("GRID", (0, 0), (-1, -1), 0.4, colors.grey),
                    ("FONTSIZE", (0, 0), (-1, -1), 7),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold")
                ]
            )
        )

        story.append(result_table)

    story.append(Spacer(1, 14))

    if chart_png:

        story.append(
            Paragraph(
                "<b>AI-Generated Chart</b>",
                styles["Heading2"]
            )
        )

        chart_buffer = BytesIO(chart_png)

        chart_image = Image(
            chart_buffer,
            width=6.2 * inch,
            height=3.5 * inch
        )

        story.append(chart_image)

        if chart_caption:
            story.append(
                Paragraph(
                    str(chart_caption),
                    styles["Italic"]
                )
            )

        story.append(Spacer(1, 14))

    story.append(
        Paragraph(
            "<b>AI-Generated Narrative</b>",
            styles["Heading2"]
        )
    )

    narrative_text = str(narrative).replace(
        "\n",
        "<br/>"
    )

    story.append(
        Paragraph(
            narrative_text,
            styles["BodyText"]
        )
    )

    document.build(story)

    buffer.seek(0)

    return buffer.getvalue()


def create_word_report(
    metadata,
    filters,
    question,
    result_df,
    narrative,
    chart_caption="",
    chart_png=None
):
    """
    Create a formatted Word report in memory.
    """

    document = Document()

    document.add_heading(
        "AI Analytics Report",
        level=0
    )

    document.add_heading(
        "Dataset Metadata",
        level=1
    )

    metadata_table = document.add_table(
        rows=1,
        cols=2
    )

    metadata_table.style = "Table Grid"

    metadata_table.rows[0].cells[0].text = "Item"
    metadata_table.rows[0].cells[1].text = "Value"

    for key, value in metadata.items():

        cells = metadata_table.add_row().cells

        cells[0].text = str(key)
        cells[1].text = str(value)

    document.add_heading(
        "Applied Filters",
        level=1
    )

    filter_table = document.add_table(
        rows=1,
        cols=2
    )

    filter_table.style = "Table Grid"

    filter_table.rows[0].cells[0].text = "Filter"
    filter_table.rows[0].cells[1].text = "Selected Value"

    for key, value in filters.items():

        cells = filter_table.add_row().cells

        cells[0].text = str(key)
        cells[1].text = str(value)

    document.add_heading(
        "User Question",
        level=1
    )

    document.add_paragraph(
        str(question)
    )

    document.add_heading(
        "Query Result",
        level=1
    )

    display_df = result_df.head(20).copy()

    if not display_df.empty:

        table = document.add_table(
            rows=1,
            cols=len(display_df.columns)
        )

        table.style = "Table Grid"

        for index, column in enumerate(
            display_df.columns
        ):
            table.rows[0].cells[index].text = str(
                column
            )

        for _, row in display_df.iterrows():

            cells = table.add_row().cells

            for index, value in enumerate(
                row.tolist()
            ):
                cells[index].text = str(value)

    if chart_png:

        document.add_heading(
            "AI-Generated Chart",
            level=1
        )

        chart_buffer = BytesIO(chart_png)

        document.add_picture(
            chart_buffer,
            width=Inches(6.2)
        )

        if chart_caption:

            paragraph = document.add_paragraph()

            run = paragraph.add_run(
                str(chart_caption)
            )

            run.italic = True

    document.add_heading(
        "AI-Generated Narrative",
        level=1
    )

    document.add_paragraph(
        str(narrative)
    )

    buffer = BytesIO()

    document.save(buffer)

    buffer.seek(0)

    return buffer.getvalue()